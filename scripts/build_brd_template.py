#!/usr/bin/env python3
"""Build BnK BRD docxtpl templates per language using python-docx.

Reads section labels from `config/brd_labels.yaml` and emits one .docx
template per language under `templates/brd/`. Each template:
  • embeds the static BnK logo (assets/bnk_logo.png) on the cover
  • uses docxtpl Jinja2 syntax for fields filled at render time
    ({{ project_name }}, {%tr for fr ... %}, {%p for fr ... %}, etc.)
  • shares an identical Jinja schema across languages so render code is
    language-agnostic — only labels differ.

Run:
    python scripts/build_brd_template.py

Outputs:
    templates/brd/BnK_BRD_Template_v2.0_en.docx
    templates/brd/BnK_BRD_Template_v2.0_vi.docx
    templates/brd/BnK_BRD_Template_v2.0_ja.docx
    templates/brd/BnK_BRD_Template_v2.0_zh.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
ASSETS_DIR = ROOT / "assets"
LOGO_PATH = ASSETS_DIR / "bnk_logo.png"
LABELS_PATH = ROOT / "config" / "brd_labels.yaml"
OUT_DIR = ROOT / "templates" / "brd"

VERSION = "v2.0"
ACCENT_HEX = "1F4E79"
ACCENT_RGB = RGBColor(0x1F, 0x4E, 0x79)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)


# ── Low-level OOXML helpers ──────────────────────────────────────────────────

def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_toc(doc: Document) -> None:
    """Insert a TOC field that Word will render the first time the doc opens."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _shade_cell(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _style_header_row(row, fill_hex: str = ACCENT_HEX, fg: RGBColor = WHITE_RGB) -> None:
    for cell in row.cells:
        _shade_cell(cell, fill_hex)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = fg


def _set_table_header(table, headers: list[str]) -> None:
    """Set headers AND apply shaded style. Avoids losing runs after cell.text=."""
    for idx, txt in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.bold = True
        run.font.color.rgb = WHITE_RGB
        _shade_cell(cell, ACCENT_HEX)


def _add_loop_table(
    doc: Document,
    headers: list[str],
    iter_var: str,
    iterable: str,
    body_cells: list[str],
):
    """Build a docxtpl-compatible loop table with the canonical 4-row layout.

    Row 0:  header row (visible, styled)
    Row 1:  {%tr for ITER in ITERABLE %} (whole row stripped at render)
    Row 2:  body cells (repeated per iteration)
    Row 3:  {%tr endfor %} (whole row stripped)

    The {%tr for%} and {%tr endfor %} MUST live in their own rows because
    docxtpl 0.20+'s regex captures only one {%tr%} tag per row.
    """
    n_cols = len(headers)
    assert len(body_cells) == n_cols, "body_cells must match header count"

    table = doc.add_table(rows=4, cols=n_cols)
    table.style = "Light Grid Accent 1"

    _set_table_header(table, headers)
    table.rows[1].cells[0].text = "{%tr for " + iter_var + " in " + iterable + " %}"
    for i, body in enumerate(body_cells):
        table.rows[2].cells[i].text = body
    table.rows[3].cells[0].text = "{%tr endfor %}"
    return table


def _bold_run(paragraph, text: str) -> None:
    paragraph.add_run(text).bold = True


def _bullet_loop(doc: Document, var_name: str, item_var: str = "i") -> None:
    """Emit a docxtpl paragraph-level for-loop rendering each item as a bullet.

    Three paragraphs are written so the {%p for %} and {%p endfor %} tags own
    their own paragraphs (which docxtpl removes at render time), with the
    body in between.
    """
    doc.add_paragraph("{%p for " + item_var + " in " + var_name + " %}")
    doc.add_paragraph("• {{ " + item_var + " }}")
    doc.add_paragraph("{%p endfor %}")


# ── Cover & front matter ─────────────────────────────────────────────────────

def _add_cover(doc: Document, labels: dict) -> None:
    # Logo (static embed; rebuild template to change)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        p.add_run().add_picture(str(LOGO_PATH), width=Mm(40))
    else:
        r = p.add_run("[BnK Logo]")
        r.font.bold = True

    for _ in range(3):
        doc.add_paragraph()

    # Document type banner
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(labels["doc_type"])
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = ACCENT_RGB

    # Project name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("{{ project_name }}")
    r.font.size = Pt(28)
    r.font.bold = True

    # Project code
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("{{ project_code }}")
    r.font.size = Pt(14)
    r.font.italic = True

    for _ in range(8):
        doc.add_paragraph()

    # Meta block: Version / Author / Date
    for key in ("version", "author", "created_at"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bold = p.add_run(f"{labels[key]}: ")
        bold.font.bold = True
        p.add_run("{{ " + key + " }}")

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(labels["copyright"])
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_doc_info(doc: Document, labels: dict) -> None:
    doc.add_heading(labels["doc_info"], level=1)
    doc.add_heading(labels["version_history"], level=2)

    _add_loop_table(
        doc,
        headers=[labels["th_version"], labels["th_date"], labels["th_description"], labels["th_author"]],
        iter_var="v",
        iterable="version_history",
        body_cells=["{{ v.version }}", "{{ v.date }}", "{{ v.description }}", "{{ v.author }}"],
    )


def _add_toc_page(doc: Document, labels: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(labels["toc"])
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = ACCENT_RGB
    _add_toc(doc)


# ── Body sections ────────────────────────────────────────────────────────────

def _add_introduction(doc: Document, labels: dict) -> None:
    doc.add_heading(labels["s1"], level=1)
    doc.add_heading(labels["s1_1"], level=2)
    doc.add_paragraph("{{ purpose }}")

    doc.add_heading(labels["s1_2"], level=2)
    _add_loop_table(
        doc,
        headers=[labels["th_role"], labels["th_party"], labels["th_responsibility"]],
        iter_var="a",
        iterable="intended_audience",
        body_cells=["{{ a.role }}", "{{ a.party }}", "{{ a.responsibility }}"],
    )


def _add_business_context(doc: Document, labels: dict) -> None:
    doc.add_heading(labels["s2"], level=1)
    doc.add_heading(labels["s2_1"], level=2)
    doc.add_paragraph("{{ background }}")
    doc.add_heading(labels["s2_2"], level=2)
    doc.add_paragraph("{{ objectives }}")
    doc.add_heading(labels["s2_3"], level=2)
    doc.add_heading(labels["s2_3_1"], level=3)
    _bullet_loop(doc, "constraints", "c")
    doc.add_heading(labels["s2_3_2"], level=3)
    _bullet_loop(doc, "assumptions", "a")


def _add_scope(doc: Document, labels: dict) -> None:
    doc.add_heading(labels["s3"], level=1)
    doc.add_heading(labels["s3_1"], level=2)
    _bullet_loop(doc, "scope_in", "s")
    doc.add_heading(labels["s3_2"], level=2)
    _bullet_loop(doc, "scope_out", "s")


def _add_stakeholders(doc: Document, labels: dict) -> None:
    doc.add_heading(labels["s4"], level=1)
    _add_loop_table(
        doc,
        headers=[labels["th_id"], labels["th_name"], labels["th_role"], labels["th_responsibility"]],
        iter_var="s",
        iterable="stakeholders",
        body_cells=["{{ s.id }}", "{{ s.name }}", "{{ s.role }}", "{{ s.responsibility }}"],
    )


def _add_business_requirements(doc: Document, labels: dict) -> None:
    doc.add_heading(labels["s5"], level=1)

    # 5.1 Overview
    doc.add_heading(labels["s5_1"], level=2)
    _add_loop_table(
        doc,
        headers=[labels["th_id"], labels["th_name"], labels["th_priority"], labels["th_short"]],
        iter_var="fr",
        iterable="functional_requirements",
        body_cells=["{{ fr.fr_id }}", "{{ fr.name }}", "{{ fr.priority }}", "{{ fr.short_description }}"],
    )

    # 5.2 FR Detail (paragraph loop)
    doc.add_heading(labels["s5_2"], level=2)
    doc.add_paragraph("{%p for fr in functional_requirements %}")
    doc.add_heading("{{ fr.fr_id }} – {{ fr.name }}", level=3)

    p = doc.add_paragraph()
    _bold_run(p, labels["fr_description"])
    doc.add_paragraph("{{ fr.description }}")

    p = doc.add_paragraph()
    _bold_run(p, labels["fr_priority"] + ": ")
    p.add_run("{{ fr.priority }}")

    p = doc.add_paragraph()
    _bold_run(p, labels["fr_user_stories"])
    _bullet_loop(doc, "fr.user_stories", "us")

    p = doc.add_paragraph()
    _bold_run(p, labels["fr_acceptance"])
    _bullet_loop(doc, "fr.acceptance_criteria", "ac")

    p = doc.add_paragraph()
    _bold_run(p, labels["fr_interface"])
    doc.add_paragraph("{{ fr.interface_notes }}")

    doc.add_paragraph("{%p endfor %}")

    # 5.3 NFR
    doc.add_heading(labels["s5_3"], level=2)
    _add_loop_table(
        doc,
        headers=[labels["th_category"], labels["th_metric"], labels["th_target"]],
        iter_var="n",
        iterable="nfr_rows",
        body_cells=["{{ n.category }}", "{{ n.metric }}", "{{ n.target }}"],
    )

    # 5.4 Data Requirements
    doc.add_heading(labels["s5_4"], level=2)
    doc.add_paragraph("{{ data_requirements }}")

    # 5.5 Integration Requirements
    doc.add_heading(labels["s5_5"], level=2)
    _add_loop_table(
        doc,
        headers=[labels["th_system"], labels["th_direction"], labels["th_protocol"], labels["th_note"]],
        iter_var="ig",
        iterable="integrations",
        body_cells=["{{ ig.system }}", "{{ ig.direction }}", "{{ ig.protocol }}", "{{ ig.note }}"],
    )


def _add_acceptance(doc: Document, labels: dict) -> None:
    doc.add_heading(labels["s6"], level=1)
    _bullet_loop(doc, "acceptance_criteria", "a")


def _add_glossary(doc: Document, labels: dict) -> None:
    doc.add_heading(labels["s7"], level=1)

    doc.add_heading(labels["s7_1"], level=2)
    _add_loop_table(
        doc,
        headers=[labels["th_term"], labels["th_definition"]],
        iter_var="g",
        iterable="glossary",
        body_cells=["{{ g.term }}", "{{ g.definition }}"],
    )

    doc.add_heading(labels["s7_2"], level=2)
    _add_loop_table(
        doc,
        headers=[labels["th_term"], labels["th_definition"]],
        iter_var="ab",
        iterable="abbreviations",
        body_cells=["{{ ab.term }}", "{{ ab.definition }}"],
    )


def _add_appendix(doc: Document, labels: dict) -> None:
    doc.add_heading(labels["s8"], level=1)
    # Optional introductory text
    doc.add_paragraph("{{ appendix }}")
    # Named appendix items (Appendix A: ..., Appendix B: ...)
    doc.add_heading(labels["s8_items"], level=2)
    _bullet_loop(doc, "appendix_items", "ap")


# ── Build orchestration ──────────────────────────────────────────────────────

def build_template(language: str, labels: dict) -> Path:
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_cover(doc, labels)
    _add_page_break(doc)

    _add_doc_info(doc, labels)
    _add_page_break(doc)

    _add_toc_page(doc, labels)
    _add_page_break(doc)

    _add_introduction(doc, labels)
    _add_business_context(doc, labels)
    _add_scope(doc, labels)
    _add_stakeholders(doc, labels)
    _add_business_requirements(doc, labels)
    _add_acceptance(doc, labels)
    _add_glossary(doc, labels)
    _add_appendix(doc, labels)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"BnK_BRD_Template_{VERSION}_{language}.docx"
    doc.save(out)
    return out


def main() -> int:
    if not LABELS_PATH.exists():
        print(f"ERROR: labels file missing: {LABELS_PATH}", file=sys.stderr)
        return 1

    with LABELS_PATH.open(encoding="utf-8") as f:
        all_labels = yaml.safe_load(f)

    if not LOGO_PATH.exists():
        print(
            f"WARNING: logo not found at {LOGO_PATH} — using text fallback",
            file=sys.stderr,
        )

    print(f"Building BRD templates ({VERSION}) for {len(all_labels)} languages...")
    for lang, labels in all_labels.items():
        out = build_template(lang, labels)
        print(f"  ✓ {lang:>3}  →  {out.relative_to(ROOT)}  ({out.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
